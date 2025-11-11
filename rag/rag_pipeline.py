import os
import re # Dosya adlarını işlemek için 're' kütüphanesini import ediyoruz
from docx import Document
from typing import Dict, List, Any
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import numpy as np
import time

# --- Konfigürasyon Ayarları ---
DOCUMENTS_PATH = r"C:\Users\PC1\Desktop\get_agent_log\documents"
EMBEDDING_MODEL_NAME = 'intfloat/multilingual-e5-large'

# --- YENİ FONKSİYON: Makine Adı Çıkarma ---
def extract_machine_name_from_filename(filename: str) -> str:
    name = re.sub(r'\.docx$', '', filename).strip()
    if name.startswith('(') and name.endswith(')'):
        name = name[1:-1]
    return name.strip()

def timeit(func):
    """
    Bir fonksiyonun çalışma süresini ölçmek için kullanılacak decorator.
    """
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        print(f"'{func.__name__}' fonksiyonu {end_time - start_time:.2f} saniyede tamamlandı.")
        return result
    return wrapper

@timeit
def load_documents(folder_path: str) -> Dict[str, str]:
    """
    Adım 1: Belirtilen klasördeki tüm .docx dosyalarını okur.
    """
    # (Bu fonksiyonda değişiklik yok)
    print("\n--- Adım 1: Dokümanlar Yükleniyor ---")
    documents = {}
    if not os.path.isdir(folder_path):
        print(f"[HATA] Klasör bulunamadı: {folder_path}")
        return documents

    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            try:
                doc = Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                if content:
                    documents[filename] = content
                    print(f"[OK] '{filename}' yüklendi.")
            except Exception as e:
                print(f"[HATA] '{filename}' okunurken sorun oluştu: {e}")
    return documents

@timeit
def chunk_texts(documents: Dict[str, str]) -> List[Any]:
    """
    Adım 2: Yüklenmiş metinleri anlamsal parçalara (chunk) ayırır.
    --- BU FONKSİYON GÜNCELLENDİ ---
    """
    print("\n--- Adım 2: Metinler Parçalara Ayrılıyor ---")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150,
        length_function=len,
        is_separator_regex=False,
        separators=["\n## ", "\n### ", "\n\n", "\n", ". ", ", ", " ", ""]
    )
    
    all_chunks = []
    print("Makine isimleri dosya adlarından çıkarılıyor ve metadatalara ekleniyor...")
    for filename, text in documents.items():
        # EN ÖNEMLİ KISIM: Makine adını çıkarıp metadata'ya ekliyoruz
        machine_name = extract_machine_name_from_filename(filename)
        print(f"  -> '{filename}' -> Makine Adı: '{machine_name}'")
        
        chunks = text_splitter.create_documents(
            [text],
            metadatas=[{"source": filename, "machine_name": machine_name}]
        )
        all_chunks.extend(chunks)
    
    print(f"Toplam {len(documents)} dokümandan {len(all_chunks)} adet chunk oluşturuldu.")
    return all_chunks

@timeit
def generate_embeddings(chunks: List[Any], model_name: str):
    """
    Adım 3: Metin parçalarını kullanarak embedding vektörleri oluşturur.
    """
    # (Bu fonksiyonda değişiklik yok)
    print("\n--- Adım 3: Embedding Vektörleri Oluşturuluyor ---")
    chunk_texts = [chunk.page_content for chunk in chunks]
    print(f"'{model_name}' modeli yükleniyor...")
    model = SentenceTransformer(model_name)
    print("Model başarıyla yüklendi.")
    print(f"{len(chunk_texts)} adet chunk için embedding işlemi başlıyor...")
    embeddings = model.encode(chunk_texts, show_progress_bar=True)
    print("Embedding işlemi tamamlandı.")
    return embeddings

def main():
    """
    Tüm RAG hazırlık sürecini yöneten ana fonksiyon.
    """
    # Adım 1
    documents = load_documents(DOCUMENTS_PATH)
    if not documents:
        print("İşlem durduruldu: Hiç doküman yüklenemedi.")
        return

    # Adım 2
    chunks = chunk_texts(documents)
    if not chunks:
        print("İşlem durduruldu: Hiç chunk oluşturulamadı.")
        return

    # Adım 3 - DÜZELTME: Eksik olan 'EMBEDDING_MODEL_NAME' argümanı eklendi.
    embeddings = generate_embeddings(chunks, EMBEDDING_MODEL_NAME)

    # --- Sonuç Kontrolü ---
    print("\n\n--- SÜREÇ TAMAMLANDI: ÖZET ---")
    print(f"Yüklenen Doküman Sayısı : {len(documents)}")
    print(f"Oluşturulan Chunk Sayısı: {len(chunks)}")
    # İlk chunk'ın metadatasını kontrol ederek makine adının eklendiğini doğrulayalım
    if chunks:
        print(f"Örnek Chunk Metadatası  : {chunks[0].metadata}")
    print(f"Oluşturulan Vektör Sayısı : {len(embeddings)}")
    print(f"Vektör Boyutu (Dimension)   : {embeddings[0].shape[0] if len(embeddings) > 0 else 'N/A'}")
    print(f"Embedding Matris Şekli : {np.array(embeddings).shape}")
    print("---------------------------------")
    print("\nArtık elimizde 'machine_name' içeren chunk'lar ve vektörler var.")
    print("Şimdi 'setup_vectordb.py' betiğini çalıştırarak Qdrant veritabanını oluşturabilirsiniz.")

if __name__ == "__main__":
    main()


# import os
# from docx import Document
# from typing import Dict, List, Any
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from sentence_transformers import SentenceTransformer
# import numpy as np
# import time

# # --- Konfigürasyon Ayarları ---
# DOCUMENTS_PATH = r"C:\Users\PC1\Desktop\get_agent_log\documents"
# EMBEDDING_MODEL_NAME = 'intfloat/multilingual-e5-large'

# def timeit(func):
#     """
#     Bir fonksiyonun çalışma süresini ölçmek için kullanılacak decorator.
#     """
#     def wrapper(*args, **kwargs):
#         start_time = time.time()
#         result = func(*args, **kwargs)
#         end_time = time.time()
#         print(f"'{func.__name__}' fonksiyonu {end_time - start_time:.2f} saniyede tamamlandı.")
#         return result
#     return wrapper

# @timeit
# def load_documents(folder_path: str) -> Dict[str, str]:
#     """
#     Adım 1: Belirtilen klasördeki tüm .docx dosyalarını okur.
#     """
#     print("\n--- Adım 1: Dokümanlar Yükleniyor ---")
#     documents = {}
#     if not os.path.isdir(folder_path):
#         print(f"[HATA] Klasör bulunamadı: {folder_path}")
#         return documents

#     for filename in os.listdir(folder_path):
#         if filename.endswith(".docx"):
#             file_path = os.path.join(folder_path, filename)
#             try:
#                 doc = Document(file_path)
#                 content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
#                 if content:
#                     documents[filename] = content
#                     print(f"[OK] '{filename}' yüklendi.")
#             except Exception as e:
#                 print(f"[HATA] '{filename}' okunurken sorun oluştu: {e}")
#     return documents

# @timeit
# def chunk_texts(documents: Dict[str, str]) -> List[Any]:
#     """
#     Adım 2: Yüklenmiş metinleri anlamsal parçalara (chunk) ayırır.
#     """
#     print("\n--- Adım 2: Metinler Parçalara Ayrılıyor ---")
#     text_splitter = RecursiveCharacterTextSplitter(
#         chunk_size=800,           # Azaltıldı: Tablolar ve listeler için daha uygun
#         chunk_overlap=150,        # Azaltıldı: Teknik içerikte daha hassas kontrol
#         length_function=len,
#         is_separator_regex=False,
#         separators=[
#             "\n## ",              # Başlıkları önceliklendir
#             "\n### ",             # Alt başlıklar
#             "\n\n",               # Paragraf sonları
#             "\n",                 # Satır sonları
#             ". ",                 # Cümle sonları
#             ", ",                 # Virgüller
#             " ",                  # Boşluklar
#             ""                    # Son çare
#         ]
#     )
    
#     all_chunks = []
#     for filename, text in documents.items():
#         # create_documents metodu, her chunk'a kaynak dosya adını
#         # metadata olarak eklememizi sağlar.
#         chunks = text_splitter.create_documents([text], metadatas=[{"source": filename}])
#         all_chunks.extend(chunks)
    
#     print(f"Toplam {len(documents)} dokümandan {len(all_chunks)} adet chunk oluşturuldu.")
#     return all_chunks

# # YENİ, DÜZELTİLMİŞ HALİ
# @timeit
# def generate_embeddings(chunks: List[Any], model_name: str):
#     """
#     Adım 3: Metin parçalarını kullanarak embedding vektörleri oluşturur.
#     """
#     print("\n--- Adım 3: Embedding Vektörleri Oluşturuluyor ---")

#     # Sadece metin içeriklerini bir listeye alıyoruz
#     chunk_texts = [chunk.page_content for chunk in chunks]

#     # Modeli yüklüyoruz (ilk seferde indirilir)
#     # Fonksiyona parametre olarak gelen model_name'i kullanıyoruz.
#     print(f"'{model_name}' modeli yükleniyor...")
#     model = SentenceTransformer(model_name)
#     print("Model başarıyla yüklendi.")

#     # Tüm metinleri vektörlere dönüştürüyoruz
#     print(f"{len(chunk_texts)} adet chunk için embedding işlemi başlıyor...")
#     embeddings = model.encode(chunk_texts, show_progress_bar=True)
#     print("Embedding işlemi tamamlandı.")

#     return embeddings

# def main():
#     """
#     Tüm RAG hazırlık sürecini yöneten ana fonksiyon.
#     """
#     # Adım 1
#     documents = load_documents(DOCUMENTS_PATH)
#     if not documents:
#         print("İşlem durduruldu: Hiç doküman yüklenemedi.")
#         return

#     # Adım 2
#     chunks = chunk_texts(documents)
#     if not chunks:
#         print("İşlem durduruldu: Hiç chunk oluşturulamadı.")
#         return

#     # Adım 3
#     embeddings = generate_embeddings(chunks)

#     # --- Sonuç Kontrolü ---
#     print("\n\n--- SÜREÇ TAMAMLANDI: ÖZET ---")
#     print(f"Yüklenen Doküman Sayısı : {len(documents)}")
#     print(f"Oluşturulan Chunk Sayısı: {len(chunks)}")
#     print(f"Oluşturulan Vektör Sayısı : {len(embeddings)}")
#     print(f"Vektör Boyutu (Dimension)   : {embeddings[0].shape[0] if len(embeddings) > 0 else 'N/A'}")
#     print(f"Embedding Matris Şekli : {np.array(embeddings).shape}")
#     print("---------------------------------")
#     print("\nArtık elimizde chunk'lar ve onlara karşılık gelen embedding vektörleri var.")
#     print("Bir sonraki adım bu verileri Qdrant vektör veritabanına yüklemek olacak.")


# if __name__ == "__main__":
#     main()